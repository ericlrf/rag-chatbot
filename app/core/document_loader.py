from dataclasses import dataclass
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader

SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}


@dataclass(frozen=True)
class LoadedDocument:
    path: Path
    text: str
    page: int | None = None


def _read_text_file(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return path.read_text(encoding="latin-1")


def _load_pdf(path: Path) -> list[LoadedDocument]:
    reader = PdfReader(str(path))
    docs: list[LoadedDocument] = []
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            docs.append(LoadedDocument(path=path, text=text, page=index))
    return docs


def _load_docx(path: Path) -> list[LoadedDocument]:
    document = DocxDocument(str(path))
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    text = "\n\n".join(paragraphs)
    return [LoadedDocument(path=path, text=text, page=None)] if text.strip() else []


def load_document(path: Path) -> list[LoadedDocument]:
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        return []

    if suffix == ".pdf":
        return _load_pdf(path)
    if suffix == ".docx":
        return _load_docx(path)

    text = _read_text_file(path)
    return [LoadedDocument(path=path, text=text, page=None)] if text.strip() else []


def iter_document_files(directory: Path) -> list[Path]:
    if not directory.exists():
        raise FileNotFoundError(f"Diretório não encontrado: {directory}")

    files = [
        path
        for path in directory.rglob("*")
        if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS
    ]
    return sorted(files)
